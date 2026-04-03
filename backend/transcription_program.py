import os
import shutil
import tempfile
import time
from pathlib import Path
from uuid import uuid4

from docx import Document
from fastapi import UploadFile
from pydub import AudioSegment
import whisper

from .usage_tracker import MonthlyUsageTracker


class TranscriptionService:
    SUPPORTED_EXTENSIONS = {".mp3", ".wav", ".opus", ".m4a", ".aac", ".ogg", ".flac"}

    def __init__(self, model_name: str | None = None) -> None:
        self.model_name = model_name or os.getenv("WHISPER_MODEL", "tiny")
        self.base_dir = Path(__file__).resolve().parent
        self.export_dir = self.base_dir / "exports"
        self.export_dir.mkdir(exist_ok=True)
        self.model = None
        self.usage_tracker = MonthlyUsageTracker()

    async def transcribe_upload(self, file: UploadFile) -> dict[str, str | None]:
        suffix = Path(file.filename or "").suffix.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            allowed = ", ".join(sorted(self.SUPPORTED_EXTENSIONS))
            raise ValueError(f"Unsupported file type. Allowed types: {allowed}")

        temp_input_path: str | None = None
        temp_wav_path: str | None = None

        try:
            self.usage_tracker.ensure_allowed()
            temp_input_fd, temp_input_path = tempfile.mkstemp(suffix=suffix)
            os.close(temp_input_fd)
            with open(temp_input_path, "wb") as temp_input:
                temp_input.write(await file.read())

            self._ensure_ffmpeg_available()
            audio = AudioSegment.from_file(temp_input_path)
            temp_wav_fd, temp_wav_path = tempfile.mkstemp(suffix=".wav")
            os.close(temp_wav_fd)

            audio.export(temp_wav_path, format="wav")
            started_at = time.perf_counter()
            result = self._get_model().transcribe(temp_wav_path)
            self.usage_tracker.record_usage(time.perf_counter() - started_at)
            transcript_text = result.get("text", "").strip()
            export_name = self._build_export_name(file.filename or "uploaded_file")
            export_path = self.export_dir / export_name
            self._save_transcript_docx(
                export_path=export_path,
                source_filename=file.filename or "uploaded_file",
                transcript_text=transcript_text,
                detected_language=result.get("language"),
            )

            return {
                "filename": file.filename or "uploaded_file",
                "text": transcript_text,
                "detected_language": result.get("language"),
                "model": self.model_name,
                "document_filename": export_name,
                "document_download_url": f"/downloads/{export_name}",
            }
        finally:
            await file.close()

            for path in (temp_input_path, temp_wav_path):
                if path and os.path.exists(path):
                    try:
                        os.remove(path)
                    except PermissionError:
                        pass

    def _ensure_ffmpeg_available(self) -> None:
        if shutil.which("ffmpeg") and shutil.which("ffprobe"):
            return

        raise RuntimeError(
            "ffmpeg is not installed or not available on PATH. "
            "Install ffmpeg and ffprobe, then restart the API."
        )

    def _get_model(self):
        if self.model is None:
            # Lazy-load Whisper so the web server can start before model load completes.
            self.model = whisper.load_model(self.model_name)

        return self.model

    def _build_export_name(self, filename: str) -> str:
        stem = Path(filename).stem.strip() or "transcript"
        safe_stem = "".join(char if char.isalnum() or char in {"-", "_"} else "_" for char in stem)
        return f"{safe_stem}_{uuid4().hex[:8]}.docx"

    def _save_transcript_docx(
        self,
        export_path: Path,
        source_filename: str,
        transcript_text: str,
        detected_language: str | None,
    ) -> None:
        document = Document()
        document.add_heading("Transcription", level=1)
        document.add_paragraph(f"Source file: {source_filename}")
        if detected_language:
            document.add_paragraph(f"Detected language: {detected_language}")
        document.add_paragraph("")
        document.add_paragraph(transcript_text or "[No transcription text returned]")
        document.save(export_path)
